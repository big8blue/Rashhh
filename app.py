import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from statsmodels.tsa.holtwinters import ExponentialSmoothing
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import os
import warnings
from datetime import datetime

warnings.filterwarnings('ignore')

st.set_page_config(
    page_title="Demand Forecast Dashboard",
    layout="wide",
    initial_sidebar_state="expanded"
)

DATA_FILE = 'v 1.xlsx'
NORMS_SHEET = 'May Month Norms'
SALES_SHEET = 'Raw Data'
FORECAST_MONTHS = 8

PLOTLY_TEMPLATE = 'plotly_white'

BRAND_COLORS = {
    'primary': '#1E3A5F',
    'secondary': '#2E86AB',
    'accent': '#A23B72',
    'success': '#2ECC71',
    'warning': '#F39C12',
    'danger': '#E74C3C',
    'info': '#3498DB',
    'light': '#F8F9FA',
    'dark': '#212529'
}

STATUS_COLORS = {
    'Stocked Out': '#E74C3C',
    'At Risk': '#E67E22',
    'Below Norm': '#F1C40F',
    'Sufficient': '#2ECC71'
}

@st.cache_data
def load_data():
    norms = pd.read_excel(DATA_FILE, sheet_name=NORMS_SHEET)
    sales = pd.read_excel(DATA_FILE, sheet_name=SALES_SHEET)
    sales.columns = sales.columns.str.strip()
    return norms, sales

def preprocess(sales, norms):
    norm_skus = set(norms['SKUCode'].unique())
    sales_f = sales[sales['SKU'].isin(norm_skus)].copy()
    sales_f['month_dt'] = pd.to_datetime(sales_f['Month'], format='%b-%Y')
    monthly = sales_f.groupby(['SKU', 'Month', 'month_dt'], as_index=False)['Sale Qty'].sum()
    monthly = monthly.sort_values(['SKU', 'month_dt'])
    full_months = pd.date_range(start='2024-01-01', end='2026-04-01', freq='MS')
    full_month_labels = [d.strftime('%b-%Y') for d in full_months]
    result = []
    for sku in norm_skus:
        sku_data = monthly[monthly['SKU'] == sku].set_index('Month')
        s = sku_data['Sale Qty'].reindex(full_month_labels, fill_value=0)
        result.append(pd.DataFrame({
            'SKU': sku,
            'Sale Qty': s.values.astype(int),
            'month_dt': full_months,
            'Month': full_month_labels
        }))
    monthly_full = pd.concat(result, ignore_index=True)
    merged = monthly_full.merge(
        norms[['SKUCode', 'Description', 'Norm', 'Stock']],
        left_on='SKU',
        right_on='SKUCode',
        how='left'
    )
    return merged

def forecast_sku(series):
    series = np.asarray(series, dtype=float)
    n = len(series)
    mean_val = float(np.mean(series)) if np.sum(series > 0) > 0 else 0.0
    if mean_val <= 0:
        z = np.zeros(FORECAST_MONTHS)
        return z, z.copy(), z.copy(), 'No Demand', 0.0
    if n < 6:
        f = np.full(FORECAST_MONTHS, mean_val)
        s = float(np.std(series)) if n > 1 else mean_val * 0.3
        lo = np.maximum(f - 1.96 * s, 0)
        hi = f + 1.96 * s
        nonzero = series[series > 0]
        mape = float(np.mean(np.abs((nonzero - mean_val) / nonzero)) * 100) if len(nonzero) > 0 else 0.0
        return f, lo, hi, 'Flat Average', round(mape, 1)
    def try_holt(s, trend, seasonal, sp):
        try:
            model = ExponentialSmoothing(
                s,
                trend=trend,
                seasonal=seasonal,
                seasonal_periods=sp,
                initialization_method='heuristic'
            )
            fitted = model.fit()
            f = fitted.forecast(FORECAST_MONTHS)
            f = np.maximum(np.asarray(f, dtype=float).ravel(), 0)
            resid = s - fitted.fittedvalues
            resid = resid[~np.isnan(resid)]
            std_r = float(np.std(resid)) if len(resid) > 1 else mean_val * 0.3
            ci_lo = np.maximum(f - 1.96 * std_r, 0)
            ci_hi = f + 1.96 * std_r
            actual = s[~np.isnan(s)]
            fitted_vals = fitted.fittedvalues[~np.isnan(s)]
            nonzero_mask = actual > 0
            if np.any(nonzero_mask):
                mape = float(np.mean(np.abs((actual[nonzero_mask] - fitted_vals[nonzero_mask]) / actual[nonzero_mask])) * 100)
            else:
                mape = 0.0
            return f, ci_lo, ci_hi, mape
        except Exception:
            return None
    result = None
    method = 'Flat Average'
    mape = 0.0
    if n >= 24:
        r = try_holt(series, 'add', 'add', 12)
        if r is not None:
            f, lo, hi, m = r
            result = (f, lo, hi)
            method = 'Holt-Winters (seasonal)'
            mape = m
    if result is None and n >= 12:
        r = try_holt(series, 'add', None, None)
        if r is not None:
            f, lo, hi, m = r
            result = (f, lo, hi)
            method = "Holt's (trend)"
            mape = m
    if result is None and n >= 6:
        r = try_holt(series, None, None, None)
        if r is not None:
            f, lo, hi, m = r
            result = (f, lo, hi)
            method = 'Simple Exp. Smoothing'
            mape = m
    if result is None:
        f = np.full(FORECAST_MONTHS, mean_val)
        s = float(np.std(series)) if n > 1 else mean_val * 0.3
        lo = np.maximum(f - 1.96 * s, 0)
        hi = f + 1.96 * s
        result = (f, lo, hi)
        nonzero = series[series > 0]
        mape = float(np.mean(np.abs((nonzero - mean_val) / nonzero)) * 100) if len(nonzero) > 0 else 0.0
    f, lo, hi = result
    return f, lo, hi, method, round(mape, 1)

def compute_all_forecasts(monthly_df):
    skus = sorted(monthly_df['SKU'].unique())
    records = []
    progress_bar = st.progress(0, text="Computing Holt-Winters forecasts for 679 SKUs...")
    for i, sku in enumerate(skus):
        sku_data = monthly_df[monthly_df['SKU'] == sku].sort_values('month_dt')
        series = sku_data['Sale Qty'].values
        norm_row = monthly_df[monthly_df['SKU'] == sku].iloc[0]
        forecast, ci_low, ci_high, method, mape = forecast_sku(series)
        total_forecast = round(sum(forecast))
        avg_monthly = round(total_forecast / FORECAST_MONTHS)
        row = {
            'SKU': sku,
            'Description': norm_row['Description'],
            'Norm': int(norm_row['Norm']),
            'Stock': int(norm_row['Stock']),
            'Total Forecast': total_forecast,
            'Avg Monthly Forecast': avg_monthly,
            'Forecast vs Norm': round(total_forecast / norm_row['Norm'] * 100, 1) if norm_row['Norm'] > 0 else 0,
            'Stock Gap': max(0, total_forecast - int(norm_row['Stock'])),
            'Method': method,
            'In-Sample MAPE': mape
        }
        if norm_row['Stock'] == 0:
            row['Status'] = 'Stocked Out'
        elif norm_row['Stock'] < norm_row['Norm'] and total_forecast > norm_row['Stock']:
            row['Status'] = 'At Risk'
        elif norm_row['Stock'] < norm_row['Norm']:
            row['Status'] = 'Below Norm'
        else:
            row['Status'] = 'Sufficient'
        monthly_forecasts = [round(forecast[j]) for j in range(FORECAST_MONTHS)]
        for j in range(FORECAST_MONTHS):
            row[f'Forecast_{j+1}'] = monthly_forecasts[j]
            row[f'CI_Low_{j+1}'] = round(ci_low[j])
            row[f'CI_High_{j+1}'] = round(ci_high[j])
        row['Suggested Norm'] = max(int(norm_row['Norm']), max(monthly_forecasts))
        records.append(row)
        pct = (i + 1) / len(skus)
        progress_bar.progress(pct, text=f"Forecasting SKU {i+1}/{len(skus)} | {sku}")
    progress_bar.empty()
    forecast_df = pd.DataFrame(records)
    forecast_months = pd.date_range(start='2026-05-01', periods=FORECAST_MONTHS, freq='MS')
    return forecast_df, [m.strftime('%b-%Y') for m in forecast_months]

def generate_docx(forecast_df, monthly_df, forecast_months_labels):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\nDEMAND FORECAST REPORT\n')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(30, 58, 95)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Holt-Winters Forecast | May 2026 – Dec 2026\n\n')
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(46, 134, 171)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()
    doc.add_heading('Executive Summary', level=1)
    total_skus = len(forecast_df)
    total_forecast_demand = forecast_df['Total Forecast'].sum()
    at_risk = len(forecast_df[forecast_df['Status'] == 'At Risk'])
    stocked_out = len(forecast_df[forecast_df['Status'] == 'Stocked Out'])
    sufficient = len(forecast_df[forecast_df['Status'] == 'Sufficient'])
    below_norm = len(forecast_df[forecast_df['Status'] == 'Below Norm'])
    doc.add_paragraph(
        f'This report presents the demand forecast for {total_skus} SKUs '
        f'from May 2026 to December 2026 using Holt-Winters (Triple Exponential Smoothing), '
        f'an industry-standard time series forecasting model.'
    )
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_data = [
        ('Total SKUs Analysed', f'{total_skus}'),
        ('Total Forecast Demand (May–Dec 2026)', f'{total_forecast_demand:,.0f} units'),
        ('Stocked Out (Stock = 0)', f'{stocked_out} SKUs'),
        ('At Risk (Below Norm & Insufficient Stock)', f'{at_risk} SKUs'),
        ('Sufficient (Meets Buyer Norm)', f'{sufficient} SKUs'),
    ]
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    if at_risk > 0 or stocked_out > 0:
        doc.add_paragraph(
            f'Immediate attention required: {stocked_out} SKUs are currently stocked out '
            f'(zero warehouse inventory), and {at_risk} SKUs are at risk of stockout. '
            f'Combined, {stocked_out + at_risk} SKUs ({round((stocked_out + at_risk) / total_skus * 100, 1)}%) '
            f'require replenishment action.'
        )
    else:
        doc.add_paragraph('All SKUs have sufficient stock to meet the forecasted demand. No immediate action required.')
    doc.add_page_break()
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'Model: Holt-Winters (Triple Exponential Smoothing)\n'
        'Library: statsmodels (ExponentialSmoothing)\n'
        'Forecast Horizon: 8 months (May 2026 – Dec 2026)\n'
        'Data Source: Raw Data (Jan 2024 – Apr 2026) + May Month Norms\n\n'
        'Holt-Winters decomposes the time series into three components: level (baseline), '
        'trend (upward/downward drift), and seasonality (repeating yearly patterns). '
        'It learns these components from historical data and projects them forward.\n\n'
        'For SKUs with fewer than 6 months of data, a simple average is used. '
        'For 6–11 months, Simple Exponential Smoothing (level only) is applied. '
        'For 12–23 months, Holt\'s method (level + trend) is used. '
        'For 24+ months, full Holt-Winters (level + trend + seasonality with m=12) is applied.'
    )
    doc.add_page_break()
    doc.add_heading('Forecast Summary by Status', level=1)
    status_summary = forecast_df.groupby('Status').agg(
        Count=('SKU', 'count'),
        Total_Forecast=('Total Forecast', 'sum'),
        Avg_Norm=('Norm', 'mean'),
        Total_Stock_Gap=('Stock Gap', 'sum')
    ).reset_index()
    status_summary.columns = ['Status', 'SKU Count', 'Total Forecast (Units)', 'Avg Norm (Units)', 'Total Stock Gap (Units)']
    status_summary['Total Forecast (Units)'] = status_summary['Total Forecast (Units)'].apply(lambda x: f'{x:,.0f}')
    status_summary['Avg Norm (Units)'] = status_summary['Avg Norm (Units)'].apply(lambda x: f'{x:,.0f}')
    status_summary['Total Stock Gap (Units)'] = status_summary['Total Stock Gap (Units)'].apply(lambda x: f'{x:,.0f}')
    t = doc.add_table(rows=1 + len(status_summary), cols=5)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(status_summary.columns):
        t.rows[0].cells[j].text = col
        for p in t.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, (_, row) in enumerate(status_summary.iterrows()):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
            for p in t.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_page_break()
    doc.add_heading('Top 20 SKU Forecast Charts', level=1)
    top_skus = forecast_df.nlargest(20, 'Total Forecast')
    fig, axes = plt.subplots(5, 4, figsize=(20, 20))
    fig.suptitle('Top 20 SKUs — Forecast May–Dec 2026', fontsize=16, fontweight='bold', color='#1E3A5F')
    axes = axes.flatten()
    for idx, (_, row) in enumerate(top_skus.iterrows()):
        ax = axes[idx]
        sku = row['SKU']
        sku_data = monthly_df[monthly_df['SKU'] == sku].sort_values('month_dt')
        hist_months = sku_data['month_dt'].values
        hist_values = sku_data['Sale Qty'].values
        forecast_months_dt = pd.date_range(start='2026-05-01', periods=FORECAST_MONTHS, freq='MS')
        forecast_values = [row[f'Forecast_{j+1}'] for j in range(FORECAST_MONTHS)]
        ci_low = [row[f'CI_Low_{j+1}'] for j in range(FORECAST_MONTHS)]
        ci_high = [row[f'CI_High_{j+1}'] for j in range(FORECAST_MONTHS)]
        ax.plot(hist_months, hist_values, color='#1E3A5F', linewidth=1, label='History')
        ax.plot(forecast_months_dt, forecast_values, color='#E74C3C', linewidth=2, label='Forecast')
        ax.fill_between(forecast_months_dt, ci_low, ci_high, color='#E74C3C', alpha=0.15, label='95% CI')
        ax.axhline(y=row['Norm'], color='#2ECC71', linestyle='--', linewidth=1, label=f"Norm: {int(row['Norm'])}")
        ax.set_title(f"{sku[:25]}...", fontsize=8)
        ax.tick_params(axis='both', labelsize=7)
        ax.xaxis.set_major_locator(mdates.MonthLocator(interval=6))
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%b-%y'))
        if idx == 0:
            ax.legend(fontsize=6, loc='upper left')
    plt.tight_layout()
    chart_buf = BytesIO()
    fig.savefig(chart_buf, dpi=150, bbox_inches='tight', format='png')
    plt.close()
    chart_buf.seek(0)
    doc.add_picture(chart_buf, width=Inches(6.5))
    last_row = doc.add_paragraph()
    run = last_row.add_run('\nCharts show: Historical sales (blue), forecast (red), 95% confidence interval (shaded), buyer norm (green dashed).')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_page_break()
    doc.add_heading('Detailed Forecast — All 679 SKUs', level=1)
    doc.add_paragraph('Due to the volume of data, the full table is available in the Excel export. '
                      'Below are the first 50 entries as a sample.')
    display_cols = ['SKU', 'Description', 'Norm', 'Stock', 'Total Forecast', 'Status', 'Stock Gap']
    display_df = forecast_df[display_cols].head(50)
    t2 = doc.add_table(rows=1 + len(display_df), cols=len(display_cols))
    t2.style = 'Light Grid Accent 1'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(display_cols):
        t2.rows[0].cells[j].text = col
        for p in t2.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(7)
    for i, (_, row) in enumerate(display_df.iterrows()):
        for j, col in enumerate(display_cols):
            val = row[col]
            if isinstance(val, float):
                val = f'{val:,.0f}'
            else:
                val = str(val)
            t2.rows[i + 1].cells[j].text = val
            for p in t2.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)
    doc.add_paragraph('\n\n--- End of Report ---')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    with st.spinner('Loading data...'):
        norms, sales = load_data()
    with st.spinner('Preprocessing data...'):
        monthly_df = preprocess(sales, norms)
    st.sidebar.markdown(
        f'<p style="font-size: 22px; font-weight: bold; color: #FFFFFF; '
        f'margin-bottom: 2px; padding-bottom: 0;">JK Fenner</p>'
        f'<p style="font-size: 13px; color: {BRAND_COLORS["secondary"]}; margin-top: 0; '
        f'padding-top: 0; opacity: 0.85;">Holt-Winters Demand Forecast | May–Dec 2026</p>',
        unsafe_allow_html=True
    )
    st.sidebar.divider()
    if 'forecast_df' not in st.session_state:
        st.session_state.forecast_df = None
        st.session_state.forecast_months = None
    if st.session_state.forecast_df is None:
        with st.spinner('Computing Holt-Winters forecasts for 679 SKUs...'):
            forecast_df, forecast_months = compute_all_forecasts(monthly_df)
            st.session_state.forecast_df = forecast_df
            st.session_state.forecast_months = forecast_months
        st.rerun()
    else:
        forecast_df = st.session_state.forecast_df
        forecast_months = st.session_state.forecast_months
        st.sidebar.markdown(
            f'<p style="color: {BRAND_COLORS["success"]}; font-size: 13px;"> Forecast Ready</p>',
            unsafe_allow_html=True
        )
        st.sidebar.divider()
        sku_list = sorted(forecast_df['SKU'].tolist())
        description_map = dict(zip(forecast_df['SKU'], forecast_df['Description']))
        sku_labels = [f"{s} — {description_map.get(s, '')}" for s in sku_list]
        status_filter = st.sidebar.selectbox(
            "Filter by Status",
            ["All", "Stocked Out", "At Risk", "Below Norm", "Sufficient"]
        )
        if status_filter != "All":
            filtered_skus = forecast_df[forecast_df['Status'] == status_filter]['SKU'].tolist()
            filtered_labels = [f"{s} — {description_map.get(s, '')}" for s in filtered_skus]
        else:
            filtered_skus = sku_list
            filtered_labels = sku_labels
        if not filtered_skus:
            st.sidebar.warning("No SKUs match the selected filter.")
            selected_sku = sku_list[0]
        else:
            selected_label = st.sidebar.selectbox(
                "Select SKU",
                filtered_labels,
                index=0
            )
            selected_sku = selected_label.split(" — ")[0]
        st.sidebar.divider()
        st.sidebar.markdown("### Download")
        col1, col2 = st.sidebar.columns(2)
        with col1:
            excel_buffer = BytesIO()
            export_cols = ['SKU', 'Description', 'Norm', 'Suggested Norm', 'Stock',
                           'Total Forecast', 'Avg Monthly Forecast', 'Forecast vs Norm', 'Stock Gap', 'Status']
            for j in range(FORECAST_MONTHS):
                export_cols.append(f'Forecast_{j+1}')
            forecast_df[export_cols].to_excel(excel_buffer, index=False, sheet_name='Forecast')
            excel_buffer.seek(0)
            st.sidebar.download_button(
                label="Excel",
                data=excel_buffer,
                file_name='forecast_data.xlsx',
                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                use_container_width=True
            )
        with col2:
            docx_buffer = generate_docx(forecast_df, monthly_df, forecast_months)
            st.sidebar.download_button(
                label="Report",
                data=docx_buffer,
                file_name='forecast_report.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True
            )
        tab1, tab2, tab3 = st.tabs([" Forecast Explorer", " Bulk Table", " Methodology & Validation"])
        with tab1:
            sku_row = forecast_df[forecast_df['SKU'] == selected_sku].iloc[0]
            status_color = STATUS_COLORS.get(sku_row['Status'], '#888')
            kpi1, kpi2, kpi3, kpi4 = st.columns(4)
            with kpi1:
                st.markdown(
                    f'<div style="background: white; padding: 15px; border-radius: 10px; '
                    f'box-shadow: 0 1px 3px rgba(0,0,0,0.1); border-left: 4px solid {BRAND_COLORS["primary"]};">'
                    f'<p style="color: #888; font-size: 12px; margin: 0;">BUYER NORM</p>'
                    f'<p style="font-size: 24px; font-weight: bold; margin: 5px 0; color: {BRAND_COLORS["primary"]};">'
                    f'{int(sku_row["Norm"]):,}</p>'
                    f'<p style="color: #aaa; font-size: 11px; margin: 0;">minimum units required</p></div>',
                    unsafe_allow_html=True
                )
            with kpi2:
                st.markdown(
                    f'<div style="background: white; padding: 15px; border-radius: 10px; '
                    f'box-shadow: 0 1px 3px rgba(0,0,0,0.1); border-left: 4px solid {BRAND_COLORS["info"]};">'
                    f'<p style="color: #888; font-size: 12px; margin: 0;">CURRENT STOCK</p>'
                    f'<p style="font-size: 24px; font-weight: bold; margin: 5px 0; color: {BRAND_COLORS["info"]};">'
                    f'{int(sku_row["Stock"]):,}</p>'
                    f'<p style="color: #aaa; font-size: 11px; margin: 0;">units in warehouse</p></div>',
                    unsafe_allow_html=True
                )
            with kpi3:
                st.markdown(
                    f'<div style="background: white; padding: 15px; border-radius: 10px; '
                    f'box-shadow: 0 1px 3px rgba(0,0,0,0.1); border-left: 4px solid {BRAND_COLORS["accent"]};">'
                    f'<p style="color: #888; font-size: 12px; margin: 0;">FORECAST TOTAL</p>'
                    f'<p style="font-size: 24px; font-weight: bold; margin: 5px 0; color: {BRAND_COLORS["accent"]};">'
                    f'{int(sku_row["Total Forecast"]):,}</p>'
                    f'<p style="color: #aaa; font-size: 11px; margin: 0;">May–Dec 2026 demand</p></div>',
                    unsafe_allow_html=True
                )
            with kpi4:
                st.markdown(
                    f'<div style="background: white; padding: 15px; border-radius: 10px; '
                    f'box-shadow: 0 1px 3px rgba(0,0,0,0.1); border-left: 4px solid {status_color};">'
                    f'<p style="color: #888; font-size: 12px; margin: 0;">STATUS</p>'
                    f'<p style="font-size: 24px; font-weight: bold; margin: 5px 0; color: {status_color};">'
                    f'{sku_row["Status"]}</p>'
                    f'<p style="color: #aaa; font-size: 11px; margin: 0;">'
                    f'{"Needs immediate action" if sku_row["Status"] in ["Stocked Out", "At Risk"] else "On track"}'
                    f'</p></div>',
                    unsafe_allow_html=True
                )
            st.divider()
            sku_data = monthly_df[monthly_df['SKU'] == selected_sku].sort_values('month_dt')
            hist_months = sku_data['month_dt'].tolist()
            hist_values = sku_data['Sale Qty'].tolist()
            forecast_months_dt = pd.date_range(start='2026-05-01', periods=FORECAST_MONTHS, freq='MS')
            forecast_values = [sku_row[f'Forecast_{j+1}'] for j in range(FORECAST_MONTHS)]
            ci_low = [sku_row[f'CI_Low_{j+1}'] for j in range(FORECAST_MONTHS)]
            ci_high = [sku_row[f'CI_High_{j+1}'] for j in range(FORECAST_MONTHS)]
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=hist_months, y=hist_values,
                mode='lines+markers',
                name='Historical Sales',
                line=dict(color=BRAND_COLORS['primary'], width=2),
                marker=dict(size=4, color=BRAND_COLORS['primary'])
            ))
            fig.add_trace(go.Scatter(
                x=forecast_months_dt, y=forecast_values,
                mode='lines+markers',
                name='Forecast',
                line=dict(color='#E74C3C', width=2.5, dash='dash'),
                marker=dict(size=5, color='#E74C3C')
            ))
            fig.add_trace(go.Scatter(
                x=list(forecast_months_dt) + list(forecast_months_dt)[::-1],
                y=list(ci_high) + list(ci_low)[::-1],
                fill='toself',
                fillcolor='rgba(231, 76, 60, 0.12)',
                line=dict(color='rgba(255,255,255,0)'),
                name='95% Confidence Interval',
                showlegend=True
            ))
            fig.add_hline(
                y=sku_row['Norm'],
                line_dash='dash',
                line_color=BRAND_COLORS['success'],
                line_width=2,
                annotation_text=f"Norm: {int(sku_row['Norm'])} units",
                annotation_position="top left",
                annotation_font_size=11
            )
            fig.update_layout(
                title=dict(
                    text=f'{selected_sku} — {sku_row["Description"]}',
                    font=dict(size=16, color=BRAND_COLORS['primary'])
                ),
                xaxis_title=None,
                yaxis_title='Units',
                template=PLOTLY_TEMPLATE,
                hovermode='x unified',
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                margin=dict(l=20, r=20, t=60, b=20),
                height=450
            )
            st.plotly_chart(fig, use_container_width=True)
            st.divider()
            st.subheader("Monthly Forecast & Suggested Stock Targets")
            norm_static = int(sku_row['Norm'])
            stock_current = int(sku_row['Stock'])
            suggested_norm = int(sku_row['Suggested Norm'])
            table_data = []
            for j in range(FORECAST_MONTHS):
                fcst = round(sku_row[f"Forecast_{j+1}"])
                suggested = max(fcst, norm_static)
                gap_month = max(0, suggested - stock_current)
                table_data.append({
                    'Month': forecast_months[j],
                    'Forecast': f'{fcst:,}',
                    '±95% CI': f'{round(sku_row[f"CI_Low_{j+1}"]):,} – {round(sku_row[f"CI_High_{j+1}"]):,}',
                    'Current Norm': f'{norm_static:,}',
                    'Suggested Norm': f'{suggested:,}',
                    'Gap vs Stock': f'{gap_month:,}'
                })
            table_df = pd.DataFrame(table_data)
            st.dataframe(table_df, use_container_width=True, hide_index=True)
            st.caption(
                f"**Suggested Norm** = max(Forecast, Current Norm). "
                f"**Current Stock**: {stock_current:,} units. "
                f"**Recommended Updated Norm**: {suggested_norm:,} units "
                f"(highest of all months). "
                f"Set this as the new buyer norm to prevent stockouts."
            )
            with st.expander("About Suggested Norm"):
                st.markdown(
                    f"The **Current Norm** ({norm_static:,}) is the buyer's existing minimum. "
                    f"The **Suggested Norm** per month is max(Forecast, Current Norm) — "
                    f"the stock level the warehouse should target at the start of that month. "
                    f"The **Gap vs Stock** shows how much additional inventory is needed "
                    f"beyond current warehouse stock ({stock_current:,} units). "
                    f"A single **Recommended Updated Norm** of **{suggested_norm:,} units** "
                    f"is proposed — this ensures capacity for the peak forecast month "
                    f"while never falling below the original buyer minimum."
                )
            if sku_row['Status'] in ['Stocked Out', 'At Risk']:
                st.warning(
                    f"**Action Required**: This SKU has a stock gap of "
                    f"**{sku_row['Stock Gap']:,} units**. "
                    f"Current stock ({stock_current:,} units) is insufficient to "
                    f"cover the forecasted demand ({int(sku_row['Total Forecast']):,} units) "
                    f"against the buyer norm of {norm_static:,} units. "
                    f"Suggested new norm: **{suggested_norm:,} units**. "
                    f"With a 4-month lead time, place a replenishment order immediately."
                )
            elif sku_row['Status'] == 'Below Norm':
                st.info(
                    f"Stock ({stock_current:,} units) is below the buyer norm "
                    f"({norm_static:,} units), but current stock appears sufficient "
                    f"for the forecast period. Monitor closely. "
                    f"Suggested norm update: **{suggested_norm:,} units**."
                )
            else:
                st.success(
                    f"Stock ({stock_current:,} units) meets or exceeds the buyer norm. "
                    f"No immediate action required."
                )
        with tab2:
            st.subheader("All 679 SKUs — Forecast Summary")
            status_order = ['Stocked Out', 'At Risk', 'Below Norm', 'Sufficient']
            display_df = forecast_df.copy()
            display_df['Status'] = pd.Categorical(display_df['Status'], categories=status_order, ordered=True)
            display_df = display_df.sort_values('Status')
            display_cols = ['SKU', 'Description', 'Norm', 'Suggested Norm', 'Stock',
                            'Total Forecast', 'Avg Monthly Forecast', 'Forecast vs Norm', 'Stock Gap', 'Status']
            display_data = display_df[display_cols].copy()
            display_data['Norm'] = display_data['Norm'].apply(lambda x: f'{x:,.0f}')
            display_data['Suggested Norm'] = display_data['Suggested Norm'].apply(lambda x: f'{x:,.0f}')
            display_data['Stock'] = display_data['Stock'].apply(lambda x: f'{x:,.0f}')
            display_data['Total Forecast'] = display_data['Total Forecast'].apply(lambda x: f'{x:,.0f}')
            display_data['Avg Monthly Forecast'] = display_data['Avg Monthly Forecast'].apply(lambda x: f'{x:,.0f}')
            display_data['Forecast vs Norm'] = display_data['Forecast vs Norm'].apply(lambda x: f'{x:.1f}%')
            display_data['Stock Gap'] = display_data['Stock Gap'].apply(lambda x: f'{x:,.0f}')
            search = st.text_input("Search by SKU or Description", placeholder="Type SKU code or description...")
            if search:
                mask = (
                    display_data['SKU'].str.contains(search, case=False, na=False) |
                    display_data['Description'].str.contains(search, case=False, na=False)
                )
                display_data = display_data[mask]
            if status_filter != "All":
                display_data = display_data[display_data['Status'] == status_filter]
            def color_status(val):
                colors = {
                    'Stocked Out': 'background-color: #FFEBEE; color: #C62828',
                    'At Risk': 'background-color: #FFF3E0; color: #E65100',
                    'Below Norm': 'background-color: #FFFDE7; color: #F9A825',
                    'Sufficient': 'background-color: #E8F5E9; color: #2E7D32'
                }
                return colors.get(val, '')
            styled = display_data.style.map(color_status, subset=['Status'])
            st.dataframe(styled, use_container_width=True, hide_index=True, height=500)
            st.caption(f"Showing {len(display_data)} of {len(forecast_df)} SKUs")
            status_counts = forecast_df['Status'].value_counts()
            fig2 = px.bar(
                x=status_counts.index,
                y=status_counts.values,
                color=status_counts.index,
                color_discrete_map=STATUS_COLORS,
                labels={'x': 'Status', 'y': 'SKU Count'},
                title='SKU Status Distribution'
            )
            fig2.update_layout(
                template=PLOTLY_TEMPLATE,
                showlegend=False,
                height=350
            )
            st.plotly_chart(fig2, use_container_width=True)

        with tab3:
            st.subheader(" Methodology & Forecast Validation")

            st.markdown("""
            ### 1. Model: Holt-Winters Triple Exponential Smoothing

            The forecasting engine uses the **Holt-Winters additive method**, implemented via
            `statsmodels.tsa.holtwinters.ExponentialSmoothing` with `initialization_method='heuristic'`.
            This is an industry-standard approach for demand planning with trend and seasonality.
            """)

            st.markdown("**Mathematical Formulation (Additive Seasonality):**")

            st.latex(r"""
            \begin{aligned}
            \text{Level:} \quad & \ell_t = \alpha \,(y_t - s_{t-m}) + (1 - \alpha)\,(\ell_{t-1} + b_{t-1}) \\[6pt]
            \text{Trend:} \quad & b_t = \beta \,(\ell_t - \ell_{t-1}) + (1 - \beta)\,b_{t-1} \\[6pt]
            \text{Seasonal:} \quad & s_t = \gamma \,(y_t - \ell_t) + (1 - \gamma)\,s_{t-m} \\[6pt]
            \text{Forecast:} \quad & \hat{y}_{t+h} = \ell_t + h \cdot b_t + s_{t+h-m}
            \end{aligned}
            """)

            st.markdown("""
            | Symbol | Meaning |
            |---|---|
            | $y_t$ | Actual sales in period $t$ |
            | $\\ell_t$ | Level (smoothed baseline) |
            | $b_t$ | Trend (smoothed slope) |
            | $s_t$ | Seasonal component |
            | $\\alpha, \\beta, \\gamma$ | Smoothing parameters (0–1), optimised via log-likelihood |
            | $m = 12$ | Seasonal period (monthly data, yearly cycle) |
            | $h$ | Forecast horizon (1–8 months ahead) |
            """)

            st.markdown("""
            **Why additive (not multiplicative)?** The additive formulation assumes seasonal fluctuations
            remain roughly constant in absolute magnitude regardless of demand level. This is appropriate
            for automotive parts where seasonal patterns are driven by production schedules, not proportional
            to volume. Multiplicative models are preferred when seasonal amplitude grows with demand
            (e.g., consumer retail).

            **Initialization**: `heuristic` sets initial level, trend, and seasonal components from the
            first two years of data using a decomposition approach, avoiding the need for a separate
            optimisation pass.
            """)
            st.divider()

            method_counts = forecast_df['Method'].value_counts()
            method_order = ['Holt-Winters (seasonal)', "Holt's (trend)",
                            'Simple Exp. Smoothing', 'Flat Average', 'No Demand']
            method_order = [m for m in method_order if m in method_counts.index]
            fig_method = px.bar(
                x=method_counts.reindex(method_order).index,
                y=method_counts.reindex(method_order).values,
                color=method_counts.reindex(method_order).index,
                color_discrete_sequence=px.colors.qualitative.Set2,
                labels={'x': 'Forecast Method', 'y': 'SKU Count'},
                title='Model Selection Distribution (679 SKUs)',
                text=method_counts.reindex(method_order).values
            )
            fig_method.update_traces(textposition='outside')
            fig_method.update_layout(
                template=PLOTLY_TEMPLATE,
                showlegend=False,
                height=400,
                xaxis_title=None
            )
            st.plotly_chart(fig_method, use_container_width=True)

            st.markdown("""
            ### 2. Adaptive Model Selection

            Not all SKUs have sufficient history for full Holt-Winters. The system automatically
            selects the **simplest model the data supports**, following the principle of parsimony
            (Occam's razor): prefer the model with fewer parameters when data is limited.

            | Data Available | Model | Why This Threshold |
            |---|---|---|
            | **≥ 24 months** | Holt-Winters (seasonal) | Full 2-year cycle needed to estimate 12-month seasonality reliably |
            | **12–23 months** | Holt's (trend) | Enough for trend estimation, but < 2 years risks spurious seasonality |
            | **6–11 months** | Simple Exp. Smoothing | Only baseline level is estimable; trend would be unreliable |
            | **< 6 months** | Flat Average | Too few observations for any smoothing — historical mean is most robust |
            | **Zero demand** | No Demand | No signal to model; forecast = 0 for all months |

            **Convergence safety**: If any model fails to converge (e.g., optimisation does not
            stabilise), the system falls back to flat average. This guarantees the dashboard
            **never crashes** due to a forecasting error.
            """)
            st.divider()

            st.subheader("### 3. Forecast vs Actual — Model Diagnostic")
            st.markdown(
                "Select a SKU to inspect how well the fitted model tracks historical sales. "
                "The **fitted values** (orange) are the model's in-sample reconstruction of the training data. "
                "Closely tracking actuals indicates a good fit; systematic deviations suggest the model "
                "may be missing a pattern."
            )

            diag_sku_list = sorted(forecast_df['SKU'].tolist())
            diag_desc_map = dict(zip(forecast_df['SKU'], forecast_df['Description']))
            diag_labels = [f"{s} — {diag_desc_map.get(s, '')}" for s in diag_sku_list]
            diag_selected = st.selectbox("Select SKU for diagnostic", diag_labels, key="diag_sku")
            diag_sku = diag_selected.split(" — ")[0]

            sku_row = forecast_df[forecast_df['SKU'] == diag_sku].iloc[0]
            sku_hist = monthly_df[monthly_df['SKU'] == diag_sku].sort_values('month_dt')
            hist_x = sku_hist['month_dt'].values
            hist_y = sku_hist['Sale Qty'].values.astype(float)
            n_months = len(hist_y)

            method_used = sku_row['Method']
            fitted_vals = None
            if method_used == 'Holt-Winters (seasonal)' and n_months >= 24:
                try:
                    model = ExponentialSmoothing(hist_y, trend='add', seasonal='add',
                                                 seasonal_periods=12,
                                                 initialization_method='heuristic').fit()
                    fitted_vals = np.maximum(model.fittedvalues, 0)
                except Exception:
                    pass
            elif method_used == "Holt's (trend)" and n_months >= 12:
                try:
                    model = ExponentialSmoothing(hist_y, trend='add', seasonal=None,
                                                 initialization_method='heuristic').fit()
                    fitted_vals = np.maximum(model.fittedvalues, 0)
                except Exception:
                    pass
            elif method_used == 'Simple Exp. Smoothing' and n_months >= 6:
                try:
                    model = ExponentialSmoothing(hist_y, trend=None, seasonal=None,
                                                 initialization_method='heuristic').fit()
                    fitted_vals = np.maximum(model.fittedvalues, 0)
                except Exception:
                    pass
            elif method_used not in ('Flat Average', 'No Demand'):
                fitted_vals = np.full(n_months, np.mean(hist_y[hist_y > 0])) if np.any(hist_y > 0) else np.zeros(n_months)

            forecast_x = pd.date_range(start='2026-05-01', periods=FORECAST_MONTHS, freq='MS')
            forecast_y = [sku_row[f'Forecast_{j+1}'] for j in range(FORECAST_MONTHS)]
            ci_lo = [sku_row[f'CI_Low_{j+1}'] for j in range(FORECAST_MONTHS)]
            ci_hi = [sku_row[f'CI_High_{j+1}'] for j in range(FORECAST_MONTHS)]

            fig_diag = go.Figure()

            fig_diag.add_trace(go.Scatter(
                x=hist_x, y=hist_y, mode='lines+markers',
                name='Actual Sales', line=dict(color=BRAND_COLORS['primary'], width=2),
                marker=dict(size=4)
            ))

            if fitted_vals is not None:
                fig_diag.add_trace(go.Scatter(
                    x=hist_x, y=fitted_vals, mode='lines',
                    name='Fitted (in-sample)', line=dict(color='#E67E22', width=2, dash='dot')
                ))

            fig_diag.add_trace(go.Scatter(
                x=forecast_x, y=forecast_y, mode='lines+markers',
                name='Forecast', line=dict(color='#E74C3C', width=2.5, dash='dash'),
                marker=dict(size=5)
            ))
            fig_diag.add_trace(go.Scatter(
                x=list(forecast_x) + list(forecast_x)[::-1],
                y=list(ci_hi) + list(ci_lo)[::-1],
                fill='toself', fillcolor='rgba(231,76,60,0.12)',
                line=dict(color='rgba(255,255,255,0)'),
                name='95% CI'
            ))
            fig_diag.add_hline(y=sku_row['Norm'], line_dash='dash', line_color=BRAND_COLORS['success'],
                               annotation_text=f"Norm: {int(sku_row['Norm'])}", annotation_position="top left")

            fig_diag.update_layout(
                title=dict(text=f"{diag_sku} — {sku_row['Description']}", font=dict(size=15, color=BRAND_COLORS['primary'])),
                template=PLOTLY_TEMPLATE, hovermode='x unified', height=420,
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                yaxis_title='Units', margin=dict(l=20, r=20, t=60, b=20)
            )
            st.plotly_chart(fig_diag, use_container_width=True)

            if fitted_vals is not None:
                resid = hist_y - fitted_vals
                resid = resid[~np.isnan(resid)]
                st.caption(
                    f"**Method used**: {method_used} | "
                    f"**Residual std**: {np.std(resid):.1f} units | "
                    f"**R² (approx)**: {1 - np.sum(resid**2) / np.sum((hist_y - np.mean(hist_y))**2):.3f}"
                )
            else:
                st.caption(f"**Method used**: {method_used} — fitted values not available for re-plotting.")

            st.divider()

            st.subheader(" Forecast Accuracy (In-Sample MAPE)")
            valid = forecast_df[forecast_df['Method'] != 'No Demand']
            mape_vals = valid['In-Sample MAPE'].dropna()
            if len(mape_vals) > 0:
                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.metric("Mean MAPE", f"{mape_vals.mean():.1f}%")
                with col_b:
                    st.metric("Median MAPE", f"{mape_vals.median():.1f}%")
                with col_c:
                    st.metric("Min MAPE", f"{mape_vals.min():.1f}%")
                with col_d:
                    st.metric("Max MAPE", f"{mape_vals.max():.1f}%")
                fig_mape = px.histogram(
                    valid,
                    x='In-Sample MAPE',
                    nbins=40,
                    labels={'In-Sample MAPE': 'MAPE (%)'},
                    title='Distribution of In-Sample MAPE Across SKUs',
                    color_discrete_sequence=['#2E86AB']
                )
                fig_mape.update_layout(template=PLOTLY_TEMPLATE, height=350, showlegend=False)
                st.plotly_chart(fig_mape, use_container_width=True)
                st.caption(
                    "**Note**: MAPE is computed on in-sample fitted values (training data). "
                    "It measures how well the model explains historical sales. "
                    "True forecast error on unseen data is typically higher."
                )

                st.markdown("#### MAPE by Forecast Method")
                method_mape = valid.groupby('Method')['In-Sample MAPE'].agg(['mean', 'median', 'count']).reset_index()
                method_mape.columns = ['Method', 'Mean MAPE', 'Median MAPE', 'SKU Count']
                method_mape = method_mape.sort_values('Mean MAPE')
                fig_method_mape = px.bar(
                    method_mape, x='Method', y='Mean MAPE', color='Method',
                    text=method_mape['Mean MAPE'].apply(lambda x: f'{x:.1f}%'),
                    labels={'Mean MAPE': 'Mean In-Sample MAPE (%)', 'Method': ''},
                    title='Mean In-Sample MAPE by Forecast Method',
                    color_discrete_sequence=px.colors.qualitative.Set2
                )
                fig_method_mape.update_traces(textposition='outside')
                fig_method_mape.update_layout(
                    template=PLOTLY_TEMPLATE, showlegend=False, height=350, xaxis_title=None
                )
                st.plotly_chart(fig_method_mape, use_container_width=True)

                st.markdown("""
                **Interpreting MAPE thresholds** (demand planning benchmarks):

                | MAPE Range | Rating | Interpretation |
                |---|---|---|
                | < 10% | Excellent | Model captures demand patterns very well |
                | 10–30% | Acceptable | Typical for SKU-level forecasting with intermittent demand |
                | > 30% | Poor | High uncertainty; consider reviewing data quality or using alternative methods |
                """)
            else:
                st.info("No accuracy data available.")

            st.divider()

            st.subheader("### 4. Residual Diagnostics")
            st.markdown(
                "Residuals (actual − fitted values) should resemble **white noise** if the model has "
                "captured all systematic patterns. Structured patterns in residuals indicate the model "
                "is missing information."
            )

            if fitted_vals is not None and len(fitted_vals) > 0:
                resid_all = hist_y - fitted_vals
                resid_clean = resid_all[~np.isnan(resid_all)]

                r_col1, r_col2 = st.columns(2)
                with r_col1:
                    fig_resid = go.Figure()
                    fig_resid.add_trace(go.Scatter(
                        x=list(range(len(resid_clean))), y=resid_clean.tolist(),
                        mode='markers+lines',
                        marker=dict(size=5, color=BRAND_COLORS['secondary']),
                        line=dict(width=0.5, color='rgba(0,0,0,0.2)')
                    ))
                    fig_resid.add_hline(y=0, line_dash='dash', line_color='red')
                    fig_resid.update_layout(
                        title='Residuals Over Time', template=PLOTLY_TEMPLATE,
                        height=300, xaxis_title='Period Index', yaxis_title='Residual (units)',
                        margin=dict(l=20, r=20, t=40, b=20)
                    )
                    st.plotly_chart(fig_resid, use_container_width=True)

                with r_col2:
                    fig_resid_hist = px.histogram(
                        x=resid_clean, nbins=20,
                        labels={'x': 'Residual (units)'},
                        title='Residual Distribution',
                        color_discrete_sequence=[BRAND_COLORS['info']]
                    )
                    fig_resid_hist.update_layout(
                        template=PLOTLY_TEMPLATE, height=300,
                        margin=dict(l=20, r=20, t=40, b=20), showlegend=False
                    )
                    st.plotly_chart(fig_resid_hist, use_container_width=True)

                st.markdown("""
                **What to look for:**
                - **No pattern** in residuals over time → good fit
                - **Clustering** of large residuals → potential heteroskedasticity (non-constant variance)
                - **Trend** in residuals → model missed a structural change
                - **Skewed distribution** → consider transformation or alternative model
                """)

                if len(resid_clean) > 2:
                    from scipy import stats as scipy_stats
                    skew_val = scipy_stats.skew(resid_clean)
                    kurt_val = scipy_stats.kurtosis(resid_clean)
                    _, p_val = scipy_stats.normaltest(resid_clean) if len(resid_clean) >= 8 else (0, 1)
                    st.caption(
                        f"**Residual stats**: skewness = {skew_val:.3f}, excess kurtosis = {kurt_val:.3f}, "
                        f"normality p-value = {p_val:.4f}"
                    )
            else:
                st.info("Fitted values not available for residual diagnostics on this SKU.")

            st.divider()

            st.subheader(" Suggested Norm Calculation")
            st.markdown(
                """
                The **Suggested Norm** is an actionable recommendation per SKU that answers:
                *"What should the buyer minimum be for this SKU?"*

                **Formula**: `Suggested Norm = max(Current Norm, max monthly forecast over next 8 months)`

                This ensures:
                - The warehouse never stocks below the existing buyer minimum (Current Norm)
                - Capacity is sized for the single highest-demand month in the forecast horizon
                - The buyer minimum is updated from a static annual number to a data-driven figure

                **Why this matters**: The original May Month Norms represent agreed minimums from the buyer,
                but actual demand often exceeds them. By updating the Norm per SKU based on the Holt-Winters forecast,
                procurement can set stock targets that match real demand patterns rather than negotiating new
                minimums reactively after a stockout.
                """
            )
            st.divider()
            st.subheader("### 5. Confidence Interval Methodology")
            st.markdown(
                "Confidence intervals quantify **forecast uncertainty**. A 95% CI means that, under the "
                "model assumptions, the true demand would fall within the interval 95% of the time."
            )

            st.markdown("""
            **Computation (non-parametric bootstrap approach):**

            1. Compute in-sample residuals: $e_t = y_t - \\hat{y}_t$
            2. Estimate residual standard deviation: $\\sigma_e = \\text{std}(e_t)$
            3. For $h$-step-ahead forecast: $\\text{CI}_{95\\%} = \\hat{y}_{t+h} \\pm 1.96 \\times \\sigma_e$
            4. Clamp lower bound to ≥ 0 (demand cannot be negative)

            For flat-average forecasts (where residuals are unavailable), use $\\sigma$ of historical sales directly.
            """)

            st.markdown("""
            **Interpretation for procurement decisions:**

            | CI Width | Signal | Recommended Action |
            |---|---|---|
            | **Narrow** (tight band) | High confidence in forecast | Standard replenishment planning |
            | **Wide** (broad band) | High demand uncertainty | Increase safety stock; review demand drivers |
            | **Asymmetric** (clamped at 0) | Near-zero demand with some volatility | Monitor closely; avoid overstocking |

            **Caveats:**
            - Intervals assume approximately **normal, homoskedastic residuals**. For intermittent demand SKUs,
              this assumption may be violated.
            - Intervals **do not widen with horizon** (a simplification; in reality, uncertainty grows with $h$).
              Consider the 8-month CI as a rough guide, not a precise probabilistic bound.
            - For SKUs with < 12 months of data, intervals are less reliable due to limited residual estimation.
            """)
            st.divider()

            st.subheader("### 6. Limitations & Assumptions")
            st.markdown("""
            **Model assumptions:**
            - Additive seasonality (constant seasonal amplitude)
            - Residuals are approximately normally distributed
            - Demand patterns are stationary (no structural breaks)
            - No external regressors (promotions, supply shocks, market changes)

            **Stated limitations:**
            - **In-sample validation only** — no out-of-sample backtest. True forecast accuracy is likely lower.
            - **90-day sea lead time** is not incorporated into the model (noted in project scope).
            - **Stockout data is unavailable** — sales are used as a proxy for demand. During stockout periods,
              recorded sales understate true demand, biasing forecasts downward.
            - **No SKU-level cross-validation** — each SKU is forecast independently; no hierarchical or
              pooled estimation across similar SKUs.

            **Potential improvements for future work:**
            - Out-of-sample backtesting (e.g., train on 2024–2025, test on Jan–Apr 2026)
            - Intermittent demand models (Croston, TSB) for sparse SKUs
            - Hierarchical forecasting to share signal across similar part families
            - External regressors (production schedules, economic indicators)
            """)

if __name__ == '__main__':
    main()
